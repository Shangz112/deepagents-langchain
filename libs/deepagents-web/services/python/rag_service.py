import os
import shutil
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional
import json
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Set MinerU config path
os.environ['MINERU_TOOLS_CONFIG_JSON'] = str(Path(__file__).parent / "magic-pdf.json")

try:
    from pymilvus import MilvusClient
    MILVUS_AVAILABLE = True
except ImportError:
    MilvusClient = None
    MILVUS_AVAILABLE = False

try:
    import chromadb
    from chromadb.utils import embedding_functions
    CHROMA_AVAILABLE = True
except ImportError:
    chromadb = None
    CHROMA_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    fitz = None
    FITZ_AVAILABLE = False

# MinerU / Magic-PDF import check
try:
    from magic_pdf.tools.common import do_parse
    MAGIC_PDF_AVAILABLE = True
except ImportError:
    do_parse = None
    MAGIC_PDF_AVAILABLE = False

from rag_components import MilvusAdapter, ChromaAdapter, ChunkingStrategy

class RAGService:
    """
    Core Service for Retrieval Augmented Generation (RAG).
    Manages vector store initialization, document parsing, chunking, and ingestion/search operations.
    """
    
    def __init__(self, persist_path: str, api_key: Optional[str] = None, base_url: Optional[str] = None):
        """
        Initialize the RAG Service.
        
        Args:
            persist_path (str): Path to persist vector database data.
            api_key (str, optional): API Key for embedding services (if applicable).
            base_url (str, optional): Base URL for embedding services (if applicable).
        """
        self.persist_path = persist_path
        self.api_key = api_key
        self.base_url = base_url
        
        # Load Config
        config_path = Path(__file__).parent / "rag_config.json"
        if config_path.exists():
            with open(config_path, "r", encoding="utf-8") as f:
                self.config = json.load(f)
        else:
            self.config = {} # Fallbacks will apply
            logger.warning("rag_config.json not found, using default configuration.")
            
        self.chunker = ChunkingStrategy(self.config.get("chunking", {}))
        
        # Initialize Vector Store
        self.vector_store = self._init_vector_store()
        
    def _init_vector_store(self):
        """
        Initialize the vector store backend based on configuration.
        Tries Milvus first if configured, falls back to ChromaDB.
        """
        vs_config = self.config.get("vector_store", {})
        vs_type = vs_config.get("type", "chroma")
        
        # Initialize Embedding Function (Shared)
        ef = None
        if CHROMA_AVAILABLE:
            # TODO: Configure specific embedding function if needed, currently using Default
            ef = embedding_functions.DefaultEmbeddingFunction()
        
        # Milvus Initialization
        if vs_type == "milvus" and MILVUS_AVAILABLE:
            try:
                milvus_cfg = vs_config.get("milvus", {})
                db_file = os.path.join(self.persist_path, milvus_cfg.get("uri", "milvus_kb.db"))
                client = MilvusClient(uri=db_file)
                logger.info("Using Milvus Lite for Vector DB")
                return MilvusAdapter(
                    client, 
                    milvus_cfg.get("collection_name", "deepagents_kb"),
                    ef
                )
            except Exception as e:
                logger.error(f"Failed to init Milvus Lite: {e}.")
                if not vs_config.get("fallback_to_chroma", True):
                    return None
                logger.info("Falling back to Chroma.")
        
        # Chroma Initialization
        if CHROMA_AVAILABLE:
            logger.info("Using ChromaDB for Vector DB")
            chroma_cfg = vs_config.get("chroma", {})
            client = chromadb.PersistentClient(path=self.persist_path)
            return ChromaAdapter(
                client,
                chroma_cfg.get("collection_name", "deepagents_kb"),
                ef
            )
            
        logger.error("Neither Milvus nor ChromaDB available. RAG disabled.")
        return None

    def ingest_file(self, file_path: str, original_filename: str, template: str = "default", source_id: Optional[str] = None) -> dict:
        """
        Ingest a file into the knowledge base.
        
        Args:
            file_path (str): Local path to the file.
            original_filename (str): Name of the file as uploaded by user.
            template (str): Processing template to use (e.g. 'default', 'scientific_paper').
            source_id (str, optional): Unique ID for the source. If provided, used as 'source' metadata.
            
        Returns:
            dict: Result status including chunk count or error message.
        """
        if not self.vector_store:
            return {"error": "RAG service unavailable"}

        path = Path(file_path)
        text_content = ""
        
        # 1. Parse Document Strategy
        # Determine file type
        ext = path.suffix.lower().replace(".", "")
        parse_strategies = self.config.get("parsing", {}).get("strategies", {}).get(ext, ["text"])
        
        # Execute strategies in order
        for strategy in parse_strategies:
            try:
                if strategy == "mineru":
                    if MAGIC_PDF_AVAILABLE:
                        logger.info(f"Using MinerU to parse {original_filename}...")
                        text_content = self._parse_pdf_miner_u(str(path))
                        if text_content: break
                elif strategy == "pymupdf":
                    if FITZ_AVAILABLE:
                        logger.info(f"Using PyMuPDF to parse {original_filename}...")
                        text_content = self._parse_pdf_pymupdf(str(path))
                        if text_content: break
                elif strategy == "docx_text":
                    # Placeholder for python-docx
                    try:
                        import docx
                        doc = docx.Document(path)
                        full_text = []
                        
                        # Iterate through element content to preserve order of paragraphs and tables
                        # Note: python-docx doesn't provide a unified iterator easily, 
                        # so we iterate body elements using the underlying XML element
                        from docx.document import Document
                        from docx.text.paragraph import Paragraph
                        from docx.table import Table
                        
                        # Revised approach: Just iterate paragraphs. Then iterate tables.
                        # Ideally we want them in order. 
                        # Let's try a safe "all content" extraction.
                        
                        for para in doc.paragraphs:
                            # Basic structure preservation: assume styles map to headers
                            text = para.text.strip()
                            if not text:
                                continue
                                
                            if para.style.name.startswith('Heading 1'):
                                full_text.append(f"# {text}")
                            elif para.style.name.startswith('Heading 2'):
                                full_text.append(f"## {text}")
                            elif para.style.name.startswith('Heading 3'):
                                full_text.append(f"### {text}")
                            else:
                                full_text.append(text)
                        
                        # Append tables as Markdown tables
                        if doc.tables:
                            full_text.append("\n\n# Tables\n")
                            for table in doc.tables:
                                # Simple markdown table conversion
                                rows = []
                                for row in table.rows:
                                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                                    rows.append("| " + " | ".join(cells) + " |")
                                
                                if rows:
                                    full_text.append("\n".join(rows))
                                    full_text.append("") # Empty line after table

                        text_content = "\n\n".join(full_text)
                        if text_content: break
                    except ImportError:
                        logger.warning("python-docx not installed. Skipping docx strategy.")
                    except Exception as e:
                        logger.warning(f"Failed to parse DOCX: {e}")
                elif strategy == "text":
                    text_content = path.read_text(encoding='utf-8', errors='ignore')
                    if text_content: break
            except Exception as e:
                logger.warning(f"Strategy {strategy} failed for {original_filename}: {e}")
                continue

        if not text_content:
            return {"error": "No text content extracted after trying all strategies"}

        # 2. Chunking with Template
        chunks = self.chunker.split(text_content, template)
        
        if not chunks:
            return {"error": "No chunks generated"}

        # 3. Vector Store Insertion
        # Use source_id if provided, otherwise fallback to filename
        source_key = source_id if source_id else original_filename
        metadatas = [{"source": source_key, "chunk_index": i} for i in range(len(chunks))]
        try:
            self.vector_store.ingest(chunks, metadatas)
        except Exception as e:
            logger.error(f"Ingestion failed for {original_filename}: {e}")
            return {"error": f"Ingestion failed: {e}"}
        
        return {
            "status": "success",
            "chunks": len(chunks),
            "source": original_filename,
            "backend": self.vector_store.__class__.__name__
        }

    def search(self, query: str, limit: int = 5) -> List[Dict]:
        """
        Search for relevant context.
        """
        return self.vector_store.search(query, limit)

    def get_chunks(self, source_name: str) -> List[Dict]:
        return self.vector_store.get_chunks(source_name)

    def update_chunk(self, chunk_id: str, text: str) -> bool:
        return self.vector_store.update_chunk(chunk_id, text)

    def delete_file(self, source_name: str) -> bool:
        """
        Remove a source from the knowledge base.
        """
        if not self.vector_store:
            return False
        return self.vector_store.delete_source(source_name)

    def get_chunks(self, source_name: str) -> List[Dict]:
        """
        Debug method to inspect chunks for a file.
        """
        if not self.vector_store:
            return []
        return self.vector_store.get_chunks(source_name)

    def _parse_pdf_miner_u(self, path: str) -> str:
        """
        Parse PDF using Magic-PDF (MinerU).
        """
        if not MAGIC_PDF_AVAILABLE:
            raise ImportError("Magic-PDF not installed")
            
        import tempfile
        import glob
        
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                with open(path, "rb") as f:
                    pdf_bytes = f.read()
                
                file_name = Path(path).stem
                
                do_parse(
                    output_dir=temp_dir,
                    pdf_file_name=file_name,
                    pdf_bytes_or_dataset=pdf_bytes,
                    model_list=[],
                    parse_method='auto',
                    f_dump_md=True,
                    f_dump_middle_json=False,
                    f_dump_model_json=False,
                    f_dump_orig_pdf=False,
                    f_dump_content_list=False,
                    f_make_md_mode='mm_md'
                )
                
                md_files = glob.glob(os.path.join(temp_dir, "**", "*.md"), recursive=True)
                
                if not md_files:
                    raise Exception("No markdown output generated by MinerU")
                
                with open(md_files[0], "r", encoding="utf-8") as f:
                    return f.read()
                    
            except Exception as e:
                logger.error(f"MinerU parsing internal error: {e}")
                raise e

    def _parse_pdf_pymupdf(self, path: str) -> str:
        """
        Fallback PDF parser using PyMuPDF.
        """
        text = ""
        with fitz.open(path) as doc:
            for page in doc:
                text += page.get_text() + "\n"
        return text

# Global instance
_rag_service = None

def get_rag_service(api_key=None, base_url=None): 
    """
    Singleton accessor for RAGService.
    Ensures only one instance of the Vector Store connection is active.
    """
    global _rag_service
    if _rag_service is None:
        # User requested override to assemble_agents directory
        db_path = Path(__file__).parent.parent.parent / "rag_data" / "vectordb"
        db_path.mkdir(parents=True, exist_ok=True)
        _rag_service = RAGService(str(db_path), api_key, base_url)
    return _rag_service
